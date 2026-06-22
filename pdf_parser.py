"""
pdf_parser.py
--------------
Utility module responsible for extracting raw text from uploaded resume
files. Supports PDF (via PyPDF2) and DOCX (via python-docx) formats.

All functions raise descriptive exceptions on failure so the calling
Streamlit app can show a friendly error message to the user instead of
crashing.
"""

from io import BytesIO
from typing import Optional

import PyPDF2

try:
    # docx support is a bonus feature - imported lazily/optionally
    import docx  # python-docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class ResumeParsingError(Exception):
    """Custom exception raised when a resume file cannot be parsed."""
    pass


def extract_text_from_pdf(uploaded_file) -> str:
    """
    Extract raw text from an uploaded PDF file.

    Args:
        uploaded_file: A file-like object (e.g. Streamlit's UploadedFile)
                       containing PDF binary data.

    Returns:
        str: The extracted, cleaned text content of the PDF.

    Raises:
        ResumeParsingError: If the file is empty, corrupted, encrypted,
                             or contains no extractable text.
    """
    if uploaded_file is None:
        raise ResumeParsingError("No file was uploaded.")

    try:
        # Ensure we read from the start of the stream
        uploaded_file.seek(0)
        file_bytes = uploaded_file.read()

        if not file_bytes:
            raise ResumeParsingError("The uploaded PDF file is empty.")

        pdf_reader = PyPDF2.PdfReader(BytesIO(file_bytes))

        if len(pdf_reader.pages) == 0:
            raise ResumeParsingError("The PDF has no pages to read.")

        if pdf_reader.is_encrypted:
            try:
                # Attempt to decrypt with an empty password (common case)
                pdf_reader.decrypt("")
            except Exception:
                raise ResumeParsingError(
                    "The PDF is password-protected and cannot be read."
                )

        extracted_text = []
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                extracted_text.append(page_text)

        full_text = "\n".join(extracted_text).strip()

        if not full_text:
            raise ResumeParsingError(
                "No readable text could be extracted from this PDF. "
                "It may be a scanned image without OCR text."
            )

        return full_text

    except ResumeParsingError:
        raise
    except PyPDF2.errors.PdfReadError:
        raise ResumeParsingError(
            "The uploaded file is not a valid or readable PDF."
        )
    except Exception as exc:
        raise ResumeParsingError(f"Unexpected error while reading PDF: {exc}")


def extract_text_from_docx(uploaded_file) -> str:
    """
    Extract raw text from an uploaded DOCX file.

    Args:
        uploaded_file: A file-like object containing DOCX binary data.

    Returns:
        str: The extracted text content of the DOCX file.

    Raises:
        ResumeParsingError: If python-docx is unavailable or the file
                             cannot be parsed.
    """
    if not DOCX_AVAILABLE:
        raise ResumeParsingError(
            "DOCX support is not installed. Run: pip install python-docx"
        )

    if uploaded_file is None:
        raise ResumeParsingError("No file was uploaded.")

    try:
        uploaded_file.seek(0)
        file_bytes = uploaded_file.read()

        if not file_bytes:
            raise ResumeParsingError("The uploaded DOCX file is empty.")

        document = docx.Document(BytesIO(file_bytes))
        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]

        # Also capture text inside tables (common in resumes)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)

        full_text = "\n".join(paragraphs).strip()

        if not full_text:
            raise ResumeParsingError(
                "No readable text could be extracted from this DOCX file."
            )

        return full_text

    except ResumeParsingError:
        raise
    except Exception as exc:
        raise ResumeParsingError(f"Unexpected error while reading DOCX: {exc}")


def extract_resume_text(uploaded_file, file_extension: Optional[str] = None) -> str:
    """
    Dispatch helper that extracts text based on the file extension.

    Args:
        uploaded_file: The uploaded file object.
        file_extension: Optional explicit extension override (e.g. 'pdf').
                         If not provided, it is inferred from the file name.

    Returns:
        str: Extracted text content.

    Raises:
        ResumeParsingError: For unsupported formats or parsing failures.
    """
    if file_extension is None:
        file_extension = uploaded_file.name.split(".")[-1].lower()

    file_extension = file_extension.lower().strip(".")

    if file_extension == "pdf":
        return extract_text_from_pdf(uploaded_file)
    elif file_extension == "docx":
        return extract_text_from_docx(uploaded_file)
    else:
        raise ResumeParsingError(
            f"Unsupported file format: .{file_extension}. "
            "Please upload a PDF or DOCX file."
        )
